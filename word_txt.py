
from docx import Document
from docx.shared import Inches

# pip3 install python-docx
# 创建word文档对象
document1 = Document()
document = Document("book/docxs/择天记.docx")
# 添加标题
document1.add_heading('择天记', 0)

# 前一行
lastLine = ''
count = 0
nextLine = ''
readNextLine = False

all_paragraphs = document.paragraphs
for paragraph in all_paragraphs:
    line = paragraph.text
    if line in '\n' \
            or '.com' in line \
            or 'http:' in line \
            or '----' in line:
        continue

    if readNextLine == True:
        nextLine = line
        readNextLine = False
        continue

    wrod = "一声响"
    if wrod in line:
        linnn = line.split(wrod)
        newLine = linnn[0] + '__' + wrod + '__' + linnn[1]
        print('下一行：', nextLine)
        print('-------')
        print('上一行：', lastLine)
        print('拆分后：', newLine)
        # 记录下一行
        readNextLine = True
        # print('当前行：', paragraph.text)
        # document1.add_paragraph(paragraph.text)

    lastLine = line
    count += 1
    if count == 100000:
        break

# document1.save("book/docxs/择天季笔录.docx")