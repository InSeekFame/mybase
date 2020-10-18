
# 安装PySide2 : pip install -i https://pypi.tuna.tsinghua.edu.cn/simple pyside2
from PySide2.QtWidgets import QApplication, QMessageBox, QLineEdit, QAction,QPlainTextEdit,QTreeWidget,QTreeWidgetItem
from PySide2.QtUiTools import QUiLoader
from PySide2.QtCore import QFile
from PySide2.QtGui import QIcon

# word
from docx import Document
from docx.shared import Inches

class MyBase:
    def __init__(self):
        # 从文件中加载UI定义
        qfile_base = QFile('ui/mybase.ui')
        qfile_base.open(QFile.ReadOnly)
        qfile_base.close()

        # 从UI定义中动态创建一个相应的窗口对象
        # 注意：里面的控件对象也成为窗口对象的属性了
        # 比如self.ui.button, self.ui.textEdit
        self.ui = QUiLoader().load(qfile_base)
        self.ui.searchBtn1.clicked.connect(self.searchBtnClick)
        self.ui.saveWordBtn.clicked.connect(self.saveWordsClick)
        self.ui.saveTextBtn.clicked.connect(self.saveTextBtnClick)
        self.ui.insertWordBtn.clicked.connect(self.insertWordBtnClick)

        # self.ui.treeWidget_top.setHeaderLabels(['情绪2'])

        # openAction = QAction()
        # openAction.setShortcut('Ctrl+o')
        # openAction.triggered.connect(self.exit)
        # self.ui.searchBtn1.addAction(openAction)

    # 录入/修改名词集录
    def insertWordBtnClick(self):
        print('录入/修改名词集录')
        isEdit = False
        document = Document("book/docxs/mybase/小说名词集录.docx")
        for paragraph in document.paragraphs:
            line = paragraph.text
            keyWords = self.ui.lineEdit.text()
            if keyWords in line \
                    and '【' in line \
                    and '】' in line:
                isEdit = True
            if '---end--' in line and isEdit == True:
                print('修改成功...')
                paragraph.insert_paragraph_before(self.ui.lineEdit2.text())
                document.save('book/docxs/mybase/小说名词集录.docx')
                isEdit = False
                break

    def saveTextBtnClick(self):
        print('------录入句录------')
        print(self.ui.textEdit1.toPlainText())
        temp_Document = Document("book/docxs/mybase/temp.docx")
        new_key_word = ''
        has_this_word = False
        for line in temp_Document.paragraphs:
            if '【' in line.text and '】' in line.text:
                print(line.text)
                new_key_word = line.text
                break

        keyWords = '【' + self.ui.lineEdit.text() + '】'
        document = Document("book/docxs/mybase/小说例句集录.docx")
        print('总行数：', document.paragraphs)
        for paragraph in document.paragraphs:
            if new_key_word in paragraph.text:
                print('收录关键词：'+paragraph.text+'==搜索关键词：'+new_key_word)
                has_this_word = True
                break
            if '---名词列表end---' in paragraph.text:
                paragraph.insert_paragraph_before(new_key_word)

        if has_this_word:
            print('已存在，不支持重复录入...')
            return
        else:
            # 添加新的词汇例句
            print('# 添加新的词汇例句')
            for paragraph in temp_Document.paragraphs:
                print(paragraph.text)

                if '【' in paragraph.text:
                    document.add_heading(paragraph.text, 3)
                elif '__' in paragraph.text:
                    # document.add_paragraph(paragraph.text, style='Intense Quote')
                    document.add_paragraph(paragraph.text)
                else:
                    document.add_paragraph(paragraph.text)

            document.save('book/docxs/mybase/小说例句集录.docx')


    def exit(self):
        print('ctrl+o')
        f = open("book/txt/搜神记[www.xiashu.cc].txt")  # 返回一个文件对象
        line = f.readline()  # 调用文件的 readline()方法
        while line:

            # print(line, end = '')　      # 在 Python 3 中使用
            line = f.readline()
            keyWords = self.ui.lineEdit.text()
            if keyWords in line:
                print(line)  # 后面跟 ',' 将忽略换行符
                self.ui.textEdit3.insertPlainText(line+'\n')
                pass
        f.close()

    def saveTextClick(self):
        print('\n\n save text')
        document = Document('book/docxs/mybase/小说例句集录.docx')
        text = self.ui.textEdit3.toPlainText() + '\n'
        document.add_paragraph(text)
        document.save('book/docxs/mybase/小说例句集录.docx')

    def saveWordsClick(self):
        print('------修改后的名词解释------')
        return
        print(self.ui.textEdit1.toPlainText())
        temp_Document = Document("book/docxs/mybase/temp.docx")
        new_key_word = ''
        has_this_word = False
        for line in temp_Document.paragraphs:
            if '【' in line.text and '】' in line.text:
                print(line.text)
                new_key_word = line.text
                break
        # return
        keyWords = '【' + self.ui.lineEdit.text() + '】'
        isEdit = False
        # # 在最末行添加
        # document.add_paragraph(self.ui.textEdit1.toPlainText())
        # document.save('book/docxs/mybase/小说名词集录.docx')
        document = Document("book/docxs/mybase/小说例句集录.docx")
        print('总行数：', document.paragraphs)
        for paragraph in document.paragraphs:
            if new_key_word in paragraph.text:
                print('%@==%@', paragraph.text)
                print('%@==%@', new_key_word)
                has_this_word = True
                break
        if has_this_word:
            print('已存在，不支持重复录入...')
            return
        else:
            # 添加新的词汇例句
            print('# 添加新的词汇例句')
            for paragraph in temp_Document.paragraphs:
                print(paragraph.text)
                document.add_paragraph(paragraph.text)
                document.save('book/docxs/mybase/小说例句集录.docx')
        return

        for index in range(len(document.paragraphs)):
            paragraph = document.paragraphs[index]
            line = paragraph.text
            keyWords = '【' + self.ui.lineEdit.text() + '】'
            if keyWords in line:
                isEdit = True
            if '---end--' in line and isEdit == True:
                # paragraph.insert_paragraph_before('aaaaaaaaa----aaaaaaa')
                last_paragraph = document.paragraphs[index-1]
                lastText = last_paragraph.text
                last_paragraph.clear()
                newWord = self.ui.lineEdit2.text()
                last_paragraph.add_run(lastText+'\n'+newWord)
                document.save('book/docxs/mybase/小说名词集录.docx')
                isEdit = False
                break

    def searchWords(self):
        print('搜索名词解释：')
        self.ui.textEdit1.clear()
        document = Document("book/docxs/mybase/小说名词集录.docx")
        isOpen = False
        # keyWords = '【' + self.ui.lineEdit.text() + '】'
        keyWords = self.ui.lineEdit.text()
        for paragraph in document.paragraphs:
            line = paragraph.text
            if keyWords in line \
                    and '【' in line \
                    and '】' in line:
                isOpen = True

            if isOpen == True:
                self.ui.textEdit1.insertPlainText(line)
                self.ui.textEdit1.insertPlainText('\n')

            if '---end--' in line and isOpen == True:
                isOpen = False


    def searchBtnClick(self):
        keyWords = self.ui.lineEdit.text()
        print(keyWords)
        # 搜索名词
        self.searchWords()
        # 搜索句子
        self.searchParagraphs()

    # 从关键字获取句子
    def searchParagraphs(self):
        print('清除内容')
        self.ui.textEdit3.clear()

        # pip3 install python-docx
        # 创建word文档对象
        document1 = Document()  # 临时目录
        wordText = '【' + self.ui.lineEdit.text() + '】'
        document1.add_heading(wordText, 3)

        document = Document("book/docxs/择天记.docx")
        # 添加标题
        # document1.add_heading('择天记', 0)

        # 前一行
        lastLine = ''
        count = 0
        nextLine = ''
        readNextLine = False

        keyWords = self.ui.lineEdit.text()
        keyWordList = keyWords.split(" ")
        for keyW in keyWordList:
            print("关键词："+keyW)
        # return

        all_paragraphs = document.paragraphs
        for paragraph in all_paragraphs:
            line = paragraph.text
            if line in '\n' \
                    or '.com' in line \
                    or 'http:' in line \
                    or '----' in line:
                continue

            if readNextLine == True:
                nextLine = line
                readNextLine = False
                continue

            word = keyWords
            word2 = ''
            word3 = ''
            print('gg')
            if len(keyWordList) == 1:
                word = keyWordList[0]
            elif len(keyWordList) == 2:
                word2 = keyWordList[1]
            else:
                word2 = keyWordList[1]
                word3 = keyWordList[2]

            if word in line and word2 in line and word3 in line:
                linnn = line.split(word)
                newLine = linnn[0] + '『' + word + '』' + linnn[1]
                print('下一行：', nextLine)
                # self.ui.textEdit3.insertPlainText(nextLine)
                print('-------')
                print('上一行：', lastLine)
                print('拆分后：', newLine)
                self.ui.textEdit3.insertPlainText(nextLine)
                self.ui.textEdit3.insertPlainText('\n\n')
                self.ui.textEdit3.insertPlainText(lastLine)
                self.ui.textEdit3.insertPlainText('\n')
                self.ui.textEdit3.insertPlainText(newLine)
                self.ui.textEdit3.insertPlainText('\n')
                # 记录下一行
                readNextLine = True
                # print('当前行：', paragraph.text)
                document1.add_paragraph(nextLine)
                document1.add_paragraph()
                document1.add_paragraph(lastLine)
                document1.add_paragraph(newLine, style='Intense Quote')


            lastLine = line
            count += 1
            if count == 10000:
                break

        document1.add_paragraph('---end---')
        document1.add_paragraph('\n\n')
        document1.save("book/docxs/mybase/temp.docx")

if __name__ == "__main__":
    app = QApplication([])
    mybase1 = MyBase()
    mybase1.ui.show()
    # mybase1.searchBtnClick()

    app.exec_()
